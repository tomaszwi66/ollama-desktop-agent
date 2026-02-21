#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
╔══════════════════════════════════════════════════════════════════════════════╗
║                                                                              ║
║                    ATLAS — AI Task & Automation System                       ║
║                                 v2.1.0                                       ║
║                                                                              ║
║   A fully autonomous AI agent powered by a local Ollama LLM that can:       ║
║     • create & edit text files, Word documents                               ║
║     • create & edit Excel workbooks with tables, charts, formatting          ║
║     • take screenshots (full screen or region)                               ║
║     • browse the web, fill forms, click buttons, scrape data                 ║
║     • execute Windows shell / PowerShell commands                            ║
║     • automate mouse & keyboard (GUI automation)                             ║
║     • generate matplotlib charts                                             ║
║     • plan → execute → verify every task automatically                       ║
║     • self-heal on errors with automatic retries                             ║
║                                                                              ║
║   Default model : jobautomation/OpenEuroLLM-Polish:latest (Ollama)          ║
║   Platform      : Windows 10 / 11                                            ║
║   Requirements  : Python 3.10+, Ollama running locally                       ║
║                                                                              ║
║   License: MIT                                                               ║
╚══════════════════════════════════════════════════════════════════════════════╝
"""

from __future__ import annotations

import os
import sys
import json
import time
import shutil
import subprocess
import re
import logging
import hashlib
import inspect
import webbrowser
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
from enum import Enum

# ════════════════════════════════════════════════════════════════
#  PATHS & DIRECTORIES
# ════════════════════════════════════════════════════════════════
SCRIPT_DIR = Path(__file__).parent.resolve()
WORKSPACE_DIR = SCRIPT_DIR
SCREENSHOTS_DIR = WORKSPACE_DIR / "screenshots"
SCREENSHOTS_DIR.mkdir(exist_ok=True)
LOG_DIR = WORKSPACE_DIR / "logs"
LOG_DIR.mkdir(exist_ok=True)

# Detect real Desktop path
DESKTOP_PATH = Path(os.path.expanduser("~/Desktop"))
if not DESKTOP_PATH.exists():
    DESKTOP_PATH = Path(os.path.expanduser("~/Pulpit"))
if not DESKTOP_PATH.exists():
    DESKTOP_PATH = (
        Path(os.environ.get("USERPROFILE", os.path.expanduser("~"))) / "Desktop"
    )

USER_HOME = Path(os.path.expanduser("~"))
DOCUMENTS_PATH = USER_HOME / "Documents"
DOWNLOADS_PATH = USER_HOME / "Downloads"

# ════════════════════════════════════════════════════════════════
#  LOGGING
# ════════════════════════════════════════════════════════════════
_log_file = LOG_DIR / f"atlas_{datetime.now():%Y%m%d_%H%M%S}.log"
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    handlers=[
        logging.FileHandler(_log_file, encoding="utf-8"),
        logging.StreamHandler(sys.stdout),
    ],
)
logger = logging.getLogger("ATLAS")

# ════════════════════════════════════════════════════════════════
#  SAFE IMPORT HELPER
# ════════════════════════════════════════════════════════════════

def _safe_import(module: str, pip_name: str | None = None):
    """Import a module, installing via pip if missing."""
    try:
        return __import__(module)
    except ImportError:
        pip_name = pip_name or module
        logger.warning("Installing missing package: %s", pip_name)
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", pip_name, "-q"],
        )
        return __import__(module)


# ════════════════════════════════════════════════════════════════
#  THIRD-PARTY IMPORTS (graceful degradation)
# ════════════════════════════════════════════════════════════════
ollama = _safe_import("ollama")

# -- PyAutoGUI ------------------------------------------------
try:
    import pyautogui

    pyautogui.FAILSAFE = True
    pyautogui.PAUSE = 0.3
except ImportError:
    pyautogui = None  # type: ignore[assignment]

# -- openpyxl --------------------------------------------------
try:
    import openpyxl
    from openpyxl.chart import BarChart, LineChart, PieChart, Reference
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    OPENPYXL_OK = True
except ImportError:
    OPENPYXL_OK = False

# -- matplotlib ------------------------------------------------
try:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    MATPLOTLIB_OK = True
except ImportError:
    plt = None  # type: ignore[assignment]
    MATPLOTLIB_OK = False

# -- Selenium --------------------------------------------------
try:
    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.common.keys import Keys
    from selenium.webdriver.support.ui import WebDriverWait, Select
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.webdriver.chrome.service import Service
    from selenium.webdriver.chrome.options import Options

    SELENIUM_OK = True
except ImportError:
    SELENIUM_OK = False

try:
    from webdriver_manager.chrome import ChromeDriverManager

    WDM_OK = True
except ImportError:
    WDM_OK = False

# -- Pillow ----------------------------------------------------
try:
    from PIL import Image, ImageGrab
except ImportError:
    Image = None  # type: ignore[assignment,misc]
    ImageGrab = None  # type: ignore[assignment,misc]

# -- Rich (pretty terminal) -----------------------------------
try:
    from rich.console import Console
    from rich.panel import Panel
    from rich.table import Table
    from rich.tree import Tree

    console = Console()
except ImportError:
    console = None  # type: ignore[assignment]

# -- python-docx -----------------------------------------------
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# -- BeautifulSoup ---------------------------------------------
try:
    from bs4 import BeautifulSoup

    BS4_OK = True
except ImportError:
    BS4_OK = False

import requests  # stdlib-level dependency

# ════════════════════════════════════════════════════════════════
#  DATA CLASSES & ENUMS
# ════════════════════════════════════════════════════════════════


class TaskStatus(Enum):
    PENDING = "⏳ Pending"
    IN_PROGRESS = "🔄 Running"
    COMPLETED = "✅ Done"
    FAILED = "❌ Failed"
    RETRYING = "🔁 Retrying"


@dataclass
class TaskStep:
    step_number: int
    description: str
    tool_name: str
    parameters: Dict[str, Any]
    status: TaskStatus = TaskStatus.PENDING
    result: Optional[str] = None
    error: Optional[str] = None
    retry_count: int = 0
    max_retries: int = 2


@dataclass
class TaskPlan:
    task_id: str
    original_request: str
    steps: List[TaskStep] = field(default_factory=list)
    status: TaskStatus = TaskStatus.PENDING
    created_at: datetime = field(default_factory=datetime.now)


@dataclass
class ToolResult:
    success: bool
    message: str
    data: Any = None
    files_created: List[str] = field(default_factory=list)


# ════════════════════════════════════════════════════════════════
#  PATH RESOLVER
# ════════════════════════════════════════════════════════════════


class PathResolver:
    """Translates human-friendly path references into real OS paths."""

    ALIASES: Dict[str, Path] = {}

    def __init__(self) -> None:
        self.ALIASES = {
            "desktop": DESKTOP_PATH,
            "pulpit": DESKTOP_PATH,
            "home": USER_HOME,
            "documents": DOCUMENTS_PATH,
            "dokumenty": DOCUMENTS_PATH,
            "downloads": DOWNLOADS_PATH,
            "pobrane": DOWNLOADS_PATH,
        }

    def resolve(self, path_str: str) -> Path:
        if not path_str:
            return WORKSPACE_DIR

        path_str = path_str.strip()

        # Expand ~
        if path_str.startswith("~"):
            return Path(os.path.expanduser(path_str))

        # Absolute path — return as-is
        p = Path(path_str)
        if p.is_absolute():
            return p

        # Check aliases in first segment
        normalised = path_str.replace("\\", "/")
        parts = normalised.split("/")
        first = parts[0].lower()

        if first in self.ALIASES:
            base = self.ALIASES[first]
            rest = "/".join(parts[1:])
            return base / rest if rest else base

        # Fallback — relative to workspace
        return WORKSPACE_DIR / path_str


path_resolver = PathResolver()

# ════════════════════════════════════════════════════════════════
#  LLM ENGINE (Ollama)
# ════════════════════════════════════════════════════════════════

_SYSTEM_PROMPT = f"""\
You are **ATLAS** — an advanced AI agent specialising in task automation on
a Windows desktop environment.  You operate by receiving a user request,
decomposing it into concrete steps, selecting the right tool for each step,
executing the plan, and verifying the outcome.

─── CORE PRINCIPLES ──────────────────────────────────────────────
1. **Plan first, act second.**  Always produce a short, efficient plan
   before doing anything.  Prefer the minimum number of steps needed.
2. **Use only the tools listed below.**  Never invent tool names.
3. **Use full, absolute file paths** based on the environment paths shown
   below.  Never guess paths.
4. **Be concise.**  Plans and verification notes should be brief and
   actionable.  No filler text.
5. **Self-heal.**  If a step fails, analyse the error and propose
   corrected parameters automatically.
6. **Verify.**  After executing all steps, confirm the result is correct.

─── ENVIRONMENT ──────────────────────────────────────────────────
Desktop   : {DESKTOP_PATH}
Home      : {USER_HOME}
Documents : {DOCUMENTS_PATH}
Downloads : {DOWNLOADS_PATH}
Workspace : {WORKSPACE_DIR}

─── AVAILABLE TOOLS ──────────────────────────────────────────────
FILE OPERATIONS
  create_text_file(path, content)      — create / overwrite a text file
  read_file(path)                      — read file contents
  edit_file(path, old_text, new_text)  — find & replace inside a file
  delete_file(path)                    — delete a file
  list_files(directory)                — list directory contents
  create_directory(path)               — create a directory tree
  copy_file(source, destination)       — copy a file
  move_file(source, destination)       — move / rename a file
  search_files(directory, pattern)     — recursive glob search
  append_to_file(path, content)        — append text to a file

EXCEL
  create_excel(path, data, sheet_name)
      data MUST be: {{"headers":["Col1","Col2"], "rows":[["a",1],["b",2]]}}
      Example budget: {{"headers":["Category","Amount"], "rows":[["Food",800],["Transport",300],["Bills",1200]]}}
  edit_excel(path, sheet_name, cell, value) — set a single cell
  add_excel_chart(path, chart_type, title)  — chart_type: bar | line | pie
  read_excel(path)                          — read all rows

SCREENSHOTS
  take_screenshot(filename)               — filename is ONLY the file name
                                            (e.g. "screen.png"), stored in
                                            the screenshots folder automatically
  screenshot_region(x, y, width, height, filename)

WEB / BROWSER
  open_url(url)
  web_fill_form(url, fields)              — fields: {{"selector":"value"}}
  web_click(selector)
  web_scrape(url, selector)

SHELL
  run_shell(command)                      — CMD
  run_powershell(command)                 — PowerShell
  get_system_info()

GUI AUTOMATION
  mouse_click(x, y)
  type_text(text)
  hotkey(keys)                            — e.g. "ctrl+c"
  wait_seconds(seconds)

CHARTS (matplotlib)
  create_chart(data, chart_type, title, filename)
      data as dict {{"Label":value}} or list

WORD DOCUMENTS
  create_word_document(path, content)
      content: string (markdown-like) or dict with title/sections

─── RESPONSE FORMAT ──────────────────────────────────────────────
When the user asks you to **perform a task**, respond with **only** a JSON
object — no extra commentary, no markdown fences, just raw JSON:

{{"plan":"<short description>","steps":[{{"step":1,"description":"<what>","tool":"<tool_name>","params":{{"<key>":"<value>"}}}}]}}

When the user asks a **question** that does not require tools, reply in
plain text (1-3 sentences max).

When asked to **verify** results, respond with:
{{"success":true/false,"note":"<brief assessment>"}}

─── SAFETY ───────────────────────────────────────────────────────
Never execute destructive system commands (format, mass delete system
files, shutdown).  If unsure, ask the user for confirmation.
"""


class OllamaEngine:
    """Manages all communication with the local Ollama LLM."""

    def __init__(
        self,
        model_name: str = "jobautomation/OpenEuroLLM-Polish:latest",
    ) -> None:
        self.model_name = model_name
        self.history: List[Dict[str, str]] = []
        self.system_prompt = _SYSTEM_PROMPT
        self.max_history = 12

    # ── core chat ────────────────────────────────────────────
    def chat(self, message: str, *, expect_json: bool = False) -> str:
        self.history.append({"role": "user", "content": message})
        if len(self.history) > self.max_history:
            self.history = self.history[-self.max_history :]

        messages = [{"role": "system", "content": self.system_prompt}]
        messages.extend(self.history)

        try:
            resp = ollama.chat(
                model=self.model_name,
                messages=messages,
                options={
                    "temperature": 0.1 if expect_json else 0.5,
                    "num_predict": 768 if expect_json else 1024,
                    "num_ctx": 2048,
                    "top_p": 0.85,
                    "repeat_penalty": 1.2,
                    "stop": (
                        ["\n\n\n", "Explanation:", "Note:", "```\n\n"]
                        if expect_json
                        else []
                    ),
                },
            )
            reply: str = resp["message"]["content"]
            self.history.append({"role": "assistant", "content": reply})
            return reply
        except Exception as exc:
            logger.error("Ollama error: %s", exc)
            return f"OLLAMA_ERROR: {exc}"

    # ── high-level helpers ───────────────────────────────────
    def plan_task(self, request: str) -> str:
        prompt = (
            f'Task: "{request}"\n'
            "Respond ONLY with the JSON plan. No explanation."
        )
        return self.chat(prompt, expect_json=True)

    def verify_result(self, task: str, results: List[str]) -> str:
        summary = "; ".join(results)[:400]
        prompt = (
            f'Task: "{task}"\nResults: {summary}\n'
            'Respond ONLY JSON: {"success":true/false,"note":"brief"}'
        )
        return self.chat(prompt, expect_json=True)

    def fix_params(self, tool: str, error: str, params: Dict) -> str:
        prompt = (
            f"Tool {tool} failed: {error}\n"
            f"Params: {json.dumps(params, ensure_ascii=False)}\n"
            'Fix and respond ONLY JSON: {"params":{...}}'
        )
        return self.chat(prompt, expect_json=True)

    def reset(self) -> None:
        self.history.clear()


# ════════════════════════════════════════════════════════════════
#  TOOLS — FILE OPERATIONS
# ════════════════════════════════════════════════════════════════


class FileTools:
    @staticmethod
    def create_text_file(path: str, content: str) -> ToolResult:
        try:
            fp = path_resolver.resolve(path)
            fp.parent.mkdir(parents=True, exist_ok=True)
            fp.write_text(content, encoding="utf-8")
            return ToolResult(True, f"Created {fp}", files_created=[str(fp)])
        except Exception as exc:
            return ToolResult(False, f"Error: {exc}")

    @staticmethod
    def read_file(path: str) -> ToolResult:
        try:
            fp = path_resolver.resolve(path)
            if not fp.exists():
                return ToolResult(False, f"Not found: {fp}")
            data = fp.read_text(encoding="utf-8")
            return ToolResult(True, f"Read {len(data)} chars", data=data)
        except Exception as exc:
            return ToolResult(False, f"Error: {exc}")

    @staticmethod
    def edit_file(path: str, old_text: str, new_text: str) -> ToolResult:
        try:
            fp = path_resolver.resolve(path)
            if not fp.exists():
                return ToolResult(False, f"Not found: {fp}")
            content = fp.read_text(encoding="utf-8")
            if old_text not in content:
                return ToolResult(False, "Text not found in file")
            fp.write_text(content.replace(old_text, new_text), encoding="utf-8")
            return ToolResult(True, f"Edited {fp}")
        except Exception as exc:
            return ToolResult(False, f"Error: {exc}")

    @staticmethod
    def delete_file(path: str) -> ToolResult:
        try:
            fp = path_resolver.resolve(path)
            if not fp.exists():
                return ToolResult(False, f"Not found: {fp}")
            fp.unlink()
            return ToolResult(True, f"Deleted {fp}")
        except Exception as exc:
            return ToolResult(False, f"Error: {exc}")

    @staticmethod
    def list_files(directory: str = ".") -> ToolResult:
        try:
            dp = path_resolver.resolve(directory)
            if not dp.exists():
                return ToolResult(False, f"Directory not found: {dp}")
            items = []
            for item in sorted(dp.iterdir()):
                icon = "📁" if item.is_dir() else "📄"
                sz = f" ({item.stat().st_size}B)" if item.is_file() else ""
                items.append(f"{icon} {item.name}{sz}")
            return ToolResult(True, f"{len(items)} items in {dp}", data=items)
        except Exception as exc:
            return ToolResult(False, f"Error: {exc}")

    @staticmethod
    def create_directory(path: str) -> ToolResult:
        try:
            dp = path_resolver.resolve(path)
            dp.mkdir(parents=True, exist_ok=True)
            return ToolResult(True, f"Directory created: {dp}")
        except Exception as exc:
            return ToolResult(False, f"Error: {exc}")

    @staticmethod
    def copy_file(source: str, destination: str) -> ToolResult:
        try:
            src = path_resolver.resolve(source)
            dst = path_resolver.resolve(destination)
            dst.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(src, dst)
            return ToolResult(True, f"Copied {src} → {dst}")
        except Exception as exc:
            return ToolResult(False, f"Error: {exc}")

    @staticmethod
    def move_file(source: str, destination: str) -> ToolResult:
        try:
            src = path_resolver.resolve(source)
            dst = path_resolver.resolve(destination)
            dst.parent.mkdir(parents=True, exist_ok=True)
            shutil.move(str(src), str(dst))
            return ToolResult(True, f"Moved {src} → {dst}")
        except Exception as exc:
            return ToolResult(False, f"Error: {exc}")

    @staticmethod
    def search_files(directory: str, pattern: str) -> ToolResult:
        try:
            dp = path_resolver.resolve(directory)
            found = [str(f) for f in dp.rglob(pattern)]
            return ToolResult(True, f"Found {len(found)} files", data=found)
        except Exception as exc:
            return ToolResult(False, f"Error: {exc}")

    @staticmethod
    def append_to_file(path: str, content: str) -> ToolResult:
        try:
            fp = path_resolver.resolve(path)
            fp.parent.mkdir(parents=True, exist_ok=True)
            with open(fp, "a", encoding="utf-8") as fh:
                fh.write(content)
            return ToolResult(True, f"Appended to {fp}")
        except Exception as exc:
            return ToolResult(False, f"Error: {exc}")


# ════════════════════════════════════════════════════════════════
#  TOOLS — EXCEL
# ════════════════════════════════════════════════════════════════


class ExcelTools:
    @staticmethod
    def create_excel(
        path: str,
        data: Any,
        sheet_name: str = "Sheet1",
    ) -> ToolResult:
        if not OPENPYXL_OK:
            return ToolResult(False, "openpyxl not installed")
        try:
            fp = path_resolver.resolve(path)
            fp.parent.mkdir(parents=True, exist_ok=True)

            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = sheet_name

            hdr_font = Font(bold=True, color="FFFFFF", size=11)
            hdr_fill = PatternFill(
                start_color="4472C4", end_color="4472C4", fill_type="solid"
            )
            bdr = Border(
                left=Side("thin"),
                right=Side("thin"),
                top=Side("thin"),
                bottom=Side("thin"),
            )
            alt_fill = PatternFill(
                start_color="D9E2F3", end_color="D9E2F3", fill_type="solid"
            )

            # ── parse data into (headers, rows) ─────────────
            headers: List[str] = []
            rows: List[List[Any]] = []

            if isinstance(data, dict):
                if "headers" in data and "rows" in data:
                    headers = data["headers"]
                    rows = data["rows"]
                elif all(isinstance(v, list) for v in data.values()):
                    headers = list(data.keys())
                    max_len = max(len(v) for v in data.values())
                    for i in range(max_len):
                        rows.append(
                            [
                                data[h][i] if i < len(data[h]) else ""
                                for h in headers
                            ]
                        )
                elif all(isinstance(v, dict) for v in data.values()):
                    sub_keys: list[str] = []
                    for v in data.values():
                        for k in v:
                            if k not in sub_keys:
                                sub_keys.append(k)
                    headers = ["Category"] + sub_keys
                    for mk, sd in data.items():
                        rows.append([mk] + [sd.get(k, "") for k in sub_keys])
                else:
                    headers = ["Category", "Value"]
                    rows = [[str(k), v] for k, v in data.items()]

            elif isinstance(data, list):
                if data and isinstance(data[0], dict):
                    headers = list(data[0].keys())
                    rows = [[it.get(h, "") for h in headers] for it in data]
                elif data and isinstance(data[0], list):
                    headers = [str(h) for h in data[0]]
                    rows = data[1:]
                else:
                    headers = ["#", "Value"]
                    rows = [[i + 1, v] for i, v in enumerate(data)]
            else:
                headers = ["Data"]
                rows = [[str(data)]]

            if not headers:
                return ToolResult(False, "Could not parse data")

            # ── write headers ────────────────────────────────
            for ci, h in enumerate(headers, 1):
                c = ws.cell(row=1, column=ci, value=str(h))
                c.font = hdr_font
                c.fill = hdr_fill
                c.alignment = Alignment(horizontal="center", vertical="center")
                c.border = bdr

            # ── write rows ───────────────────────────────────
            for ri, row in enumerate(rows, 2):
                for ci, val in enumerate(row, 1):
                    if isinstance(val, str):
                        try:
                            val = float(val) if "." in val else int(val)
                        except (ValueError, TypeError):
                            pass
                    c = ws.cell(row=ri, column=ci, value=val)
                    c.border = bdr
                    c.alignment = Alignment(horizontal="center")
                    if ri % 2 == 0:
                        c.fill = alt_fill

            # ── auto-width ───────────────────────────────────
            for ci in range(1, len(headers) + 1):
                mx = 0
                for ri in range(1, len(rows) + 2):
                    cv = ws.cell(row=ri, column=ci).value
                    if cv is not None:
                        mx = max(mx, len(str(cv)))
                ws.column_dimensions[get_column_letter(ci)].width = min(mx + 4, 50)

            # ── auto-filter ──────────────────────────────────
            last_col = get_column_letter(len(headers))
            ws.auto_filter.ref = f"A1:{last_col}{len(rows) + 1}"

            # ── SUM row if numeric columns exist ─────────────
            if len(headers) >= 2:
                numeric = any(
                    isinstance(r[1], (int, float))
                    for r in rows
                    if len(r) >= 2
                )
                if numeric:
                    sr = len(rows) + 2
                    ws.cell(row=sr, column=1, value="TOTAL").font = Font(bold=True)
                    ws.cell(row=sr, column=1).border = bdr
                    for ci in range(2, len(headers) + 1):
                        cl = get_column_letter(ci)
                        c = ws.cell(
                            row=sr,
                            column=ci,
                            value=f"=SUM({cl}2:{cl}{len(rows)+1})",
                        )
                        c.font = Font(bold=True)
                        c.border = bdr
                        c.alignment = Alignment(horizontal="center")

            wb.save(fp)
            return ToolResult(
                True,
                f"Excel created: {fp} ({len(rows)} rows × {len(headers)} cols)",
                files_created=[str(fp)],
            )
        except Exception as exc:
            logger.error("Excel error: %s", exc, exc_info=True)
            return ToolResult(False, f"Excel error: {exc}")

    @staticmethod
    def edit_excel(
        path: str,
        sheet_name: str | None = None,
        cell: str = "A1",
        value: Any = "",
    ) -> ToolResult:
        if not OPENPYXL_OK:
            return ToolResult(False, "openpyxl not installed")
        try:
            fp = path_resolver.resolve(path)
            if not fp.exists():
                return ToolResult(False, f"Not found: {fp}")
            wb = openpyxl.load_workbook(fp)
            ws = (
                wb[sheet_name]
                if sheet_name and sheet_name in wb.sheetnames
                else wb.active
            )
            ws[cell] = value
            wb.save(fp)
            return ToolResult(True, f"Set {cell}={value} in {fp}")
        except Exception as exc:
            return ToolResult(False, f"Error: {exc}")

    @staticmethod
    def add_excel_chart(
        path: str,
        chart_type: str = "bar",
        title: str = "Chart",
        **kwargs: Any,
    ) -> ToolResult:
        if not OPENPYXL_OK:
            return ToolResult(False, "openpyxl not installed")
        try:
            fp = path_resolver.resolve(path)
            if not fp.exists():
                return ToolResult(False, f"Not found: {fp}")

            wb = openpyxl.load_workbook(fp)
            ws = wb.active
            mr = ws.max_row
            mc = ws.max_column

            if mr < 2 or mc < 2:
                return ToolResult(False, "Not enough data for a chart")

            chart_cls = {
                "bar": BarChart,
                "line": LineChart,
                "pie": PieChart,
            }.get(chart_type.lower(), BarChart)
            chart = chart_cls()
            chart.title = title
            chart.width = 18
            chart.height = 10
            chart.style = 10

            if chart_type.lower() == "pie":
                data_ref = Reference(ws, min_col=2, min_row=1, max_row=mr)
                cats = Reference(ws, min_col=1, min_row=2, max_row=mr)
                chart.add_data(data_ref, titles_from_data=True)
                chart.set_categories(cats)
            else:
                data_ref = Reference(
                    ws, min_col=2, max_col=mc, min_row=1, max_row=mr
                )
                cats = Reference(ws, min_col=1, min_row=2, max_row=mr)
                chart.add_data(data_ref, titles_from_data=True)
                chart.set_categories(cats)

            ws.add_chart(chart, f"A{mr + 3}")
            wb.save(fp)
            return ToolResult(True, f"Chart '{title}' ({chart_type}) added to {fp}")
        except Exception as exc:
            return ToolResult(False, f"Chart error: {exc}")

    @staticmethod
    def read_excel(path: str, sheet_name: str | None = None) -> ToolResult:
        if not OPENPYXL_OK:
            return ToolResult(False, "openpyxl not installed")
        try:
            fp = path_resolver.resolve(path)
            if not fp.exists():
                return ToolResult(False, f"Not found: {fp}")
            wb = openpyxl.load_workbook(fp)
            ws = (
                wb[sheet_name]
                if sheet_name and sheet_name in wb.sheetnames
                else wb.active
            )
            data = [list(row) for row in ws.iter_rows(values_only=True)]
            return ToolResult(True, f"Read {len(data)} rows", data=data)
        except Exception as exc:
            return ToolResult(False, f"Error: {exc}")


# ════════════════════════════════════════════════════════════════
#  TOOLS — SCREENSHOTS
# ════════════════════════════════════════════════════════════════


class ScreenshotTools:
    @staticmethod
    def take_screenshot(filename: str | None = None) -> ToolResult:
        try:
            if not filename:
                filename = f"screen_{datetime.now():%Y%m%d_%H%M%S}.png"
            fname = Path(filename).name  # strip any directory component
            filepath = SCREENSHOTS_DIR / fname

            if ImageGrab:
                ImageGrab.grab().save(filepath)
            elif pyautogui:
                pyautogui.screenshot().save(filepath)
            else:
                return ToolResult(False, "No screenshot module available")

            return ToolResult(
                True, f"Screenshot saved: {filepath}", files_created=[str(filepath)]
            )
        except Exception as exc:
            return ToolResult(False, f"Screenshot error: {exc}")

    @staticmethod
    def screenshot_region(
        x: int, y: int, width: int, height: int, filename: str | None = None
    ) -> ToolResult:
        try:
            if not filename:
                filename = f"region_{datetime.now():%Y%m%d_%H%M%S}.png"
            fname = Path(filename).name
            filepath = SCREENSHOTS_DIR / fname

            if ImageGrab:
                ImageGrab.grab(bbox=(x, y, x + width, y + height)).save(filepath)
            elif pyautogui:
                pyautogui.screenshot(region=(x, y, width, height)).save(filepath)
            else:
                return ToolResult(False, "No screenshot module available")

            return ToolResult(
                True, f"Region screenshot: {filepath}", files_created=[str(filepath)]
            )
        except Exception as exc:
            return ToolResult(False, f"Error: {exc}")


# ════════════════════════════════════════════════════════════════
#  TOOLS — WEB / BROWSER
# ════════════════════════════════════════════════════════════════


class WebTools:
    def __init__(self) -> None:
        self.driver: Any = None

    def _ensure_driver(self) -> bool:
        if self.driver:
            try:
                _ = self.driver.title
                return True
            except Exception:
                self.driver = None

        if not SELENIUM_OK:
            return False

        try:
            opts = Options()
            opts.add_argument("--no-sandbox")
            opts.add_argument("--disable-dev-shm-usage")
            opts.add_argument("--disable-gpu")
            opts.add_argument("--window-size=1920,1080")
            opts.add_experimental_option("excludeSwitches", ["enable-logging"])

            if WDM_OK:
                svc = Service(ChromeDriverManager().install())
                self.driver = webdriver.Chrome(service=svc, options=opts)
            else:
                self.driver = webdriver.Chrome(options=opts)

            self.driver.implicitly_wait(10)
            return True
        except Exception as exc:
            logger.error("Chrome init error: %s", exc)
            return False

    def open_url(self, url: str) -> ToolResult:
        try:
            if not url.startswith(("http://", "https://")):
                url = "https://" + url

            if self._ensure_driver():
                self.driver.get(url)
                time.sleep(2)
                return ToolResult(True, f"Opened: {url} ({self.driver.title})")

            webbrowser.open(url)
            return ToolResult(True, f"Opened in default browser: {url}")
        except Exception as exc:
            return ToolResult(False, f"Error: {exc}")

    def web_fill_form(
        self, url: str | None = None, fields: Dict[str, str] | None = None
    ) -> ToolResult:
        try:
            if not self._ensure_driver():
                return ToolResult(False, "Browser unavailable")
            if url:
                self.driver.get(url)
                time.sleep(2)

            filled: List[str] = []
            for selector, value in (fields or {}).items():
                el = None
                for by in [By.ID, By.NAME, By.CSS_SELECTOR, By.XPATH]:
                    try:
                        el = self.driver.find_element(by, selector)
                        break
                    except Exception:
                        continue
                if el:
                    tag = el.tag_name.lower()
                    if tag == "select":
                        try:
                            Select(el).select_by_visible_text(value)
                        except Exception:
                            Select(el).select_by_value(value)
                    elif el.get_attribute("type") in ("checkbox", "radio"):
                        if value.lower() in ("true", "1", "yes") and not el.is_selected():
                            el.click()
                    else:
                        el.clear()
                        el.send_keys(value)
                    filled.append(selector)

            return ToolResult(
                True, f"Filled {len(filled)}/{len(fields or {})} fields"
            )
        except Exception as exc:
            return ToolResult(False, f"Error: {exc}")

    def web_click(self, selector: str) -> ToolResult:
        try:
            if not self._ensure_driver():
                return ToolResult(False, "Browser unavailable")
            for by in [
                By.ID,
                By.CSS_SELECTOR,
                By.XPATH,
                By.NAME,
                By.LINK_TEXT,
                By.PARTIAL_LINK_TEXT,
            ]:
                try:
                    el = WebDriverWait(self.driver, 5).until(
                        EC.element_to_be_clickable((by, selector))
                    )
                    el.click()
                    time.sleep(1)
                    return ToolResult(True, f"Clicked: {selector}")
                except Exception:
                    continue
            return ToolResult(False, f"Element not found: {selector}")
        except Exception as exc:
            return ToolResult(False, f"Error: {exc}")

    def web_scrape(
        self, url: str | None = None, selector: str | None = None
    ) -> ToolResult:
        try:
            html: str = ""
            if url:
                if self._ensure_driver():
                    self.driver.get(url)
                    time.sleep(2)
                    html = self.driver.page_source
                else:
                    resp = requests.get(
                        url, timeout=15, headers={"User-Agent": "Mozilla/5.0"}
                    )
                    html = resp.text
            elif self.driver:
                html = self.driver.page_source
            else:
                return ToolResult(False, "No URL and no open browser")

            if BS4_OK:
                soup = BeautifulSoup(html, "html.parser")
                if selector:
                    els = soup.select(selector)
                    data = [e.get_text(strip=True) for e in els]
                    return ToolResult(True, f"{len(data)} elements", data=data)
                return ToolResult(
                    True, "Page text extracted", data=soup.get_text("\n", strip=True)[:3000]
                )

            return ToolResult(True, "Raw HTML", data=html[:3000])
        except Exception as exc:
            return ToolResult(False, f"Error: {exc}")

    def close(self) -> None:
        if self.driver:
            try:
                self.driver.quit()
            except Exception:
                pass
            self.driver = None


# ════════════════════════════════════════════════════════════════
#  TOOLS — SHELL
# ════════════════════════════════════════════════════════════════


class ShellTools:
    BLOCKED = [
        "format c:",
        "del /s /q c:",
        "rd /s /q c:",
        "rm -rf /",
        "shutdown",
        "restart",
    ]

    @staticmethod
    def _safe(command: str) -> bool:
        cl = command.lower()
        return not any(b in cl for b in ShellTools.BLOCKED)

    @staticmethod
    def run_shell(command: str, timeout: int = 60) -> ToolResult:
        if not ShellTools._safe(command):
            return ToolResult(False, "Blocked for safety")
        try:
            r = subprocess.run(
                command,
                shell=True,
                capture_output=True,
                text=True,
                timeout=timeout,
                encoding="utf-8",
                errors="replace",
                cwd=str(WORKSPACE_DIR),
            )
            out = r.stdout.strip()
            err = r.stderr.strip()
            if r.returncode == 0:
                return ToolResult(True, "OK", data=out or "(no output)")
            return ToolResult(False, f"Exit {r.returncode}: {err or out}")
        except subprocess.TimeoutExpired:
            return ToolResult(False, f"Timeout ({timeout}s)")
        except Exception as exc:
            return ToolResult(False, f"Error: {exc}")

    @staticmethod
    def run_powershell(command: str, timeout: int = 60) -> ToolResult:
        if not ShellTools._safe(command):
            return ToolResult(False, "Blocked for safety")
        try:
            r = subprocess.run(
                ["powershell", "-NoProfile", "-Command", command],
                capture_output=True,
                text=True,
                timeout=timeout,
                encoding="utf-8",
                errors="replace",
                cwd=str(WORKSPACE_DIR),
            )
            out = r.stdout.strip()
            err = r.stderr.strip()
            if r.returncode == 0:
                return ToolResult(True, "OK", data=out or "(no output)")
            return ToolResult(False, f"Error: {err or out}")
        except Exception as exc:
            return ToolResult(False, f"Error: {exc}")

    @staticmethod
    def get_system_info() -> ToolResult:
        info: Dict[str, str] = {
            "platform": sys.platform,
            "python": sys.version.split()[0],
            "user": os.getlogin(),
            "desktop": str(DESKTOP_PATH),
            "workspace": str(WORKSPACE_DIR),
        }
        try:
            import psutil

            info["cpu"] = f"{psutil.cpu_count()} cores, {psutil.cpu_percent()}%"
            mem = psutil.virtual_memory()
            info["ram"] = f"{mem.total / (1024**3):.1f}GB, {mem.percent}% used"
        except ImportError:
            pass
        return ToolResult(True, "System info", data=info)


# ════════════════════════════════════════════════════════════════
#  TOOLS — GUI AUTOMATION
# ════════════════════════════════════════════════════════════════


class AutomationTools:
    @staticmethod
    def mouse_click(x: int, y: int, button: str = "left") -> ToolResult:
        if not pyautogui:
            return ToolResult(False, "pyautogui unavailable")
        try:
            pyautogui.click(x, y, button=button)
            return ToolResult(True, f"Clicked ({x},{y})")
        except Exception as exc:
            return ToolResult(False, f"Error: {exc}")

    @staticmethod
    def type_text(text: str) -> ToolResult:
        if not pyautogui:
            return ToolResult(False, "pyautogui unavailable")
        try:
            try:
                import pyperclip

                pyperclip.copy(text)
                pyautogui.hotkey("ctrl", "v")
            except ImportError:
                pyautogui.typewrite(text, interval=0.02)
            return ToolResult(True, f"Typed {len(text)} chars")
        except Exception as exc:
            return ToolResult(False, f"Error: {exc}")

    @staticmethod
    def hotkey(keys: str) -> ToolResult:
        if not pyautogui:
            return ToolResult(False, "pyautogui unavailable")
        try:
            parts = [k.strip() for k in keys.replace("+", ",").split(",")]
            pyautogui.hotkey(*parts)
            return ToolResult(True, f"Hotkey: {keys}")
        except Exception as exc:
            return ToolResult(False, f"Error: {exc}")

    @staticmethod
    def wait_seconds(seconds: float) -> ToolResult:
        time.sleep(seconds)
        return ToolResult(True, f"Waited {seconds}s")


# ════════════════════════════════════════════════════════════════
#  TOOLS — MATPLOTLIB CHARTS
# ════════════════════════════════════════════════════════════════


class ChartTools:
    @staticmethod
    def create_chart(
        data: Any,
        chart_type: str = "bar",
        title: str = "Chart",
        filename: str | None = None,
        xlabel: str = "",
        ylabel: str = "",
    ) -> ToolResult:
        if not MATPLOTLIB_OK:
            return ToolResult(False, "matplotlib unavailable")
        try:
            if not filename:
                filename = f"chart_{datetime.now():%Y%m%d_%H%M%S}.png"
            filepath = WORKSPACE_DIR / Path(filename).name

            fig, ax = plt.subplots(figsize=(10, 6))

            if isinstance(data, dict):
                labels = list(data.keys())
                values = list(data.values())
                if chart_type == "bar":
                    colours = plt.cm.viridis(  # type: ignore[attr-defined]
                        [i / max(len(labels), 1) for i in range(len(labels))]
                    )
                    bars = ax.bar(labels, values, color=colours)
                    for b, v in zip(bars, values):
                        ax.text(
                            b.get_x() + b.get_width() / 2,
                            b.get_height(),
                            str(v),
                            ha="center",
                            va="bottom",
                            fontweight="bold",
                        )
                elif chart_type == "line":
                    ax.plot(labels, values, marker="o", linewidth=2, markersize=8)
                elif chart_type == "pie":
                    ax.pie(values, labels=labels, autopct="%1.1f%%", startangle=90)
                elif chart_type == "scatter":
                    ax.scatter(labels, values, s=100)
            elif isinstance(data, list):
                if chart_type == "bar":
                    ax.bar(range(len(data)), data)
                else:
                    ax.plot(data, marker="o")

            ax.set_title(title, fontsize=14, fontweight="bold")
            if xlabel:
                ax.set_xlabel(xlabel)
            if ylabel:
                ax.set_ylabel(ylabel)
            if chart_type != "pie":
                plt.xticks(rotation=45, ha="right")

            plt.tight_layout()
            plt.savefig(filepath, dpi=150, bbox_inches="tight", facecolor="white")
            plt.close()
            return ToolResult(
                True, f"Chart saved: {filepath}", files_created=[str(filepath)]
            )
        except Exception as exc:
            plt.close("all")
            return ToolResult(False, f"Chart error: {exc}")


# ════════════════════════════════════════════════════════════════
#  TOOLS — WORD DOCUMENTS
# ════════════════════════════════════════════════════════════════


class DocumentTools:
    @staticmethod
    def create_word_document(path: str, content: Any) -> ToolResult:
        if not DOCX_OK:
            return ToolResult(False, "python-docx not installed")
        try:
            fp = path_resolver.resolve(path)
            fp.parent.mkdir(parents=True, exist_ok=True)
            doc = DocxDocument()

            if isinstance(content, str):
                for line in content.split("\n"):
                    if line.startswith("# "):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith("## "):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith("### "):
                        doc.add_heading(line[4:], level=3)
                    elif line.startswith("- "):
                        doc.add_paragraph(line[2:], style="List Bullet")
                    elif line.strip():
                        doc.add_paragraph(line)

            elif isinstance(content, dict):
                if "title" in content:
                    doc.add_heading(content["title"], level=0)
                for sec in content.get("sections", []):
                    if "heading" in sec:
                        doc.add_heading(sec["heading"], level=1)
                    if "text" in sec:
                        doc.add_paragraph(sec["text"])
                    for b in sec.get("bullets", []):
                        doc.add_paragraph(b, style="List Bullet")

            elif isinstance(content, list):
                for item in content:
                    doc.add_paragraph(str(item))

            doc.save(fp)
            return ToolResult(True, f"Word document: {fp}", files_created=[str(fp)])
        except Exception as exc:
            return ToolResult(False, f"Error: {exc}")


# ════════════════════════════════════════════════════════════════
#  TOOL MANAGER
# ════════════════════════════════════════════════════════════════


class ToolManager:
    """Registry and executor for all available tools."""

    def __init__(self) -> None:
        self._file = FileTools()
        self._excel = ExcelTools()
        self._ss = ScreenshotTools()
        self._web = WebTools()
        self._shell = ShellTools()
        self._auto = AutomationTools()
        self._chart = ChartTools()
        self._doc = DocumentTools()

        self.registry: Dict[str, Any] = {
            # files
            "create_text_file": self._file.create_text_file,
            "read_file": self._file.read_file,
            "edit_file": self._file.edit_file,
            "delete_file": self._file.delete_file,
            "list_files": self._file.list_files,
            "create_directory": self._file.create_directory,
            "copy_file": self._file.copy_file,
            "move_file": self._file.move_file,
            "search_files": self._file.search_files,
            "append_to_file": self._file.append_to_file,
            # excel
            "create_excel": self._excel.create_excel,
            "edit_excel": self._excel.edit_excel,
            "add_excel_chart": self._excel.add_excel_chart,
            "read_excel": self._excel.read_excel,
            # screenshots
            "take_screenshot": self._ss.take_screenshot,
            "screenshot_region": self._ss.screenshot_region,
            # web
            "open_url": self._web.open_url,
            "web_fill_form": self._web.web_fill_form,
            "web_click": self._web.web_click,
            "web_scrape": self._web.web_scrape,
            # shell
            "run_shell": self._shell.run_shell,
            "run_powershell": self._shell.run_powershell,
            "get_system_info": self._shell.get_system_info,
            # gui
            "mouse_click": self._auto.mouse_click,
            "type_text": self._auto.type_text,
            "hotkey": self._auto.hotkey,
            "wait_seconds": self._auto.wait_seconds,
            # charts
            "create_chart": self._chart.create_chart,
            # word
            "create_word_document": self._doc.create_word_document,
        }

    def execute(self, tool_name: str, params: Dict[str, Any]) -> ToolResult:
        func = self.registry.get(tool_name)
        if func is None:
            return ToolResult(False, f"Unknown tool: {tool_name}")

        try:
            sig = inspect.signature(func)
            valid = {k: v for k, v in params.items() if k in sig.parameters}
            return func(**valid) if valid else func(**params)
        except TypeError:
            try:
                return func(*params.values())
            except Exception as exc:
                return ToolResult(False, f"Param error for {tool_name}: {exc}")
        except Exception as exc:
            return ToolResult(False, f"{tool_name} failed: {exc}")

    def cleanup(self) -> None:
        self._web.close()


# ════════════════════════════════════════════════════════════════
#  RESPONSE PARSER
# ════════════════════════════════════════════════════════════════


class ResponseParser:
    """Extracts structured JSON from potentially messy LLM output."""

    @staticmethod
    def extract_json(text: str) -> Optional[Dict[str, Any]]:
        # 1) raw
        try:
            return json.loads(text)
        except (json.JSONDecodeError, ValueError):
            pass

        # 2) ```json ... ```
        m = re.search(r"```json\s*(.*?)\s*```", text, re.DOTALL)
        if m:
            try:
                return json.loads(m.group(1))
            except (json.JSONDecodeError, ValueError):
                pass

        # 3) ``` ... ```
        m = re.search(r"```\s*(.*?)\s*```", text, re.DOTALL)
        if m:
            try:
                return json.loads(m.group(1))
            except (json.JSONDecodeError, ValueError):
                pass

        # 4) balanced braces
        depth = 0
        start = -1
        for i, ch in enumerate(text):
            if ch == "{":
                if depth == 0:
                    start = i
                depth += 1
            elif ch == "}":
                depth -= 1
                if depth == 0 and start >= 0:
                    try:
                        return json.loads(text[start : i + 1])
                    except (json.JSONDecodeError, ValueError):
                        start = -1

        # 5) cleanup attempt
        cleaned = re.sub(r",\s*([\]}])", r"\1", text)
        cleaned = cleaned.replace("'", '"')
        try:
            return json.loads(cleaned)
        except (json.JSONDecodeError, ValueError):
            pass

        return None

    @staticmethod
    def parse_plan(llm_response: str) -> Optional[TaskPlan]:
        data = ResponseParser.extract_json(llm_response)
        if not data or "steps" not in data:
            return None

        tid = hashlib.md5(str(datetime.now()).encode()).hexdigest()[:8]
        plan = TaskPlan(task_id=tid, original_request=data.get("plan", ""))

        for sd in data["steps"]:
            plan.steps.append(
                TaskStep(
                    step_number=sd.get("step", len(plan.steps) + 1),
                    description=sd.get("description", ""),
                    tool_name=sd.get("tool", ""),
                    parameters=sd.get("params", {}),
                )
            )
        return plan


# ════════════════════════════════════════════════════════════════
#  EXECUTION ENGINE
# ════════════════════════════════════════════════════════════════


class ExecutionEngine:
    """Runs a TaskPlan step-by-step with retries and self-healing."""

    def __init__(self, tools: ToolManager, llm: OllamaEngine) -> None:
        self.tools = tools
        self.llm = llm

    def run(self, plan: TaskPlan) -> TaskPlan:
        plan.status = TaskStatus.IN_PROGRESS

        if console:
            console.print(
                Panel(
                    f"[bold cyan]{plan.original_request}[/]\n"
                    f"[dim]ID: {plan.task_id} | Steps: {len(plan.steps)}[/]",
                    title="🚀 Executing",
                    border_style="cyan",
                )
            )

        step_results: List[str] = []

        for step in plan.steps:
            step.status = TaskStatus.IN_PROGRESS

            if console:
                console.print(
                    f"\n[yellow]⚡ Step {step.step_number}:[/] {step.description}"
                )
                params_preview = json.dumps(step.parameters, ensure_ascii=False)[:80]
                console.print(f"   [dim]{step.tool_name}({params_preview})[/]")

            for attempt in range(step.max_retries + 1):
                result = self.tools.execute(step.tool_name, step.parameters)

                if result.success:
                    step.status = TaskStatus.COMPLETED
                    step.result = result.message
                    if console:
                        console.print(f"   [green]✅ {result.message}[/]")
                        if result.data:
                            console.print(f"   [dim]{str(result.data)[:150]}[/]")
                    break

                # ── failure handling ──────────────────────────
                if attempt < step.max_retries:
                    step.retry_count += 1
                    step.status = TaskStatus.RETRYING
                    if console:
                        console.print(
                            f"   [yellow]🔁 Retry {attempt + 2}: {result.message}[/]"
                        )
                    # Ask LLM to fix params
                    fix_resp = self.llm.fix_params(
                        step.tool_name, result.message, step.parameters
                    )
                    fix_data = ResponseParser.extract_json(fix_resp)
                    if fix_data and "params" in fix_data:
                        step.parameters = fix_data["params"]
                        if console:
                            console.print(
                                f"   [cyan]🔧 Fixed params: "
                                f"{json.dumps(step.parameters, ensure_ascii=False)[:80]}[/]"
                            )
                    time.sleep(0.5)
                else:
                    step.status = TaskStatus.FAILED
                    step.error = result.message
                    if console:
                        console.print(f"   [red]❌ {result.message}[/]")

            step_results.append(
                f"Step {step.step_number} ({step.tool_name}): "
                f"{step.result or step.error}"
            )

        # ── aggregate status ─────────────────────────────────
        n_ok = sum(1 for s in plan.steps if s.status == TaskStatus.COMPLETED)
        n_total = len(plan.steps)
        plan.status = TaskStatus.COMPLETED if n_ok > 0 else TaskStatus.FAILED

        # ── verification ─────────────────────────────────────
        if console:
            console.print("\n[magenta]🔍 Verifying…[/]")

        ver_raw = self.llm.verify_result(plan.original_request, step_results)
        ver = ResponseParser.extract_json(ver_raw)
        note = ver.get("note", "") if ver else ver_raw[:200]

        if console:
            colour = "green" if n_ok == n_total else ("yellow" if n_ok else "red")
            console.print(
                Panel(
                    f"[bold]{n_ok}/{n_total} steps succeeded[/]\n{note}",
                    title="📊 Result",
                    border_style=colour,
                )
            )

        return plan


# ════════════════════════════════════════════════════════════════
#  ATLAS AGENT (main class)
# ════════════════════════════════════════════════════════════════


class AtlasAgent:
    """
    Top-level agent controller.

    Flow:
      1. User types a request.
      2. LLM produces a JSON plan.
      3. Execution engine runs each step.
      4. LLM verifies the outcome.
      5. Results are reported to the user.
    """

    def __init__(
        self,
        model: str = "jobautomation/OpenEuroLLM-Polish:latest",
    ) -> None:
        self.llm = OllamaEngine(model)
        self.tools = ToolManager()
        self.engine = ExecutionEngine(self.tools, self.llm)
        self.history: List[TaskPlan] = []
        self._running = True

    # ── UI helpers ───────────────────────────────────────────

    def _banner(self) -> None:
        text = f"""\
╔═══════════════════════════════════════════════════════╗
║           🤖  ATLAS  Agent  v2.1.0                    ║
║      AI  Task  &  Automation  System                  ║
║      Model: {self.llm.model_name:<40} ║
╠═══════════════════════════════════════════════════════╣
║  /help  /tools  /status  /history  /clear  /exit      ║
║  Desktop   : {str(DESKTOP_PATH)[:40]:<40} ║
║  Workspace : {str(WORKSPACE_DIR)[:40]:<40} ║
╚═══════════════════════════════════════════════════════╝"""
        if console:
            console.print(text, style="bold cyan")
        else:
            print(text)

    def _help(self) -> None:
        h = """\
EXAMPLES
  • Create a file test.txt on the Desktop with "hello world"
  • Build an Excel budget spreadsheet with a pie chart
  • Take a screenshot
  • Open google.com
  • List files on Desktop
  • Create a sales bar chart
  • Get system info

COMMANDS
  /help     — this help
  /tools    — list available tools
  /status   — system & dependency status
  /history  — past tasks
  /clear    — clear conversation memory
  /exit     — quit"""
        if console:
            console.print(Panel(h, title="Help", border_style="blue"))
        else:
            print(h)

    def _show_tools(self) -> None:
        descriptions = {
            "create_text_file": "Create a text file",
            "read_file": "Read file contents",
            "edit_file": "Find & replace in file",
            "delete_file": "Delete a file",
            "list_files": "List directory contents",
            "create_directory": "Create directories",
            "copy_file": "Copy a file",
            "move_file": "Move / rename a file",
            "search_files": "Recursive file search",
            "append_to_file": "Append to a file",
            "create_excel": "Create Excel workbook",
            "edit_excel": "Edit Excel cell",
            "add_excel_chart": "Add chart to Excel",
            "read_excel": "Read Excel data",
            "take_screenshot": "Full-screen screenshot",
            "screenshot_region": "Region screenshot",
            "open_url": "Open URL in browser",
            "web_fill_form": "Fill a web form",
            "web_click": "Click web element",
            "web_scrape": "Scrape web page",
            "run_shell": "Run CMD command",
            "run_powershell": "Run PowerShell command",
            "get_system_info": "System information",
            "mouse_click": "Click at coordinates",
            "type_text": "Type text (keyboard)",
            "hotkey": "Send keyboard shortcut",
            "wait_seconds": "Wait / sleep",
            "create_chart": "Create matplotlib chart",
            "create_word_document": "Create Word document",
        }
        if console:
            tbl = Table(title="🧰 Available Tools")
            tbl.add_column("Tool", style="cyan", width=24)
            tbl.add_column("Description", width=40)
            for name in sorted(self.tools.registry):
                tbl.add_row(name, descriptions.get(name, ""))
            console.print(tbl)
        else:
            for name in sorted(self.tools.registry):
                print(f"  {name:<24} {descriptions.get(name, '')}")

    def _show_status(self) -> None:
        lines = [
            f"Model       : {self.llm.model_name}",
            f"Desktop     : {DESKTOP_PATH}",
            f"Workspace   : {WORKSPACE_DIR}",
            f"Tools       : {len(self.tools.registry)}",
            f"Tasks done  : {len(self.history)}",
            f"Selenium    : {'✅' if SELENIUM_OK else '❌'}",
            f"PyAutoGUI   : {'✅' if pyautogui else '❌'}",
            f"Matplotlib  : {'✅' if MATPLOTLIB_OK else '❌'}",
            f"python-docx : {'✅' if DOCX_OK else '❌'}",
            f"openpyxl    : {'✅' if OPENPYXL_OK else '❌'}",
            f"BeautifulSoup: {'✅' if BS4_OK else '❌'}",
        ]
        if console:
            console.print(Panel("\n".join(lines), title="Status", border_style="green"))
        else:
            print("\n".join(lines))

    def _show_history(self) -> None:
        if not self.history:
            print("No tasks yet.")
            return
        if console:
            tbl = Table(title="📜 Task History")
            tbl.add_column("ID", width=10)
            tbl.add_column("Request", width=42)
            tbl.add_column("Status", width=12)
            for t in self.history[-10:]:
                req = (
                    t.original_request[:40] + "…"
                    if len(t.original_request) > 40
                    else t.original_request
                )
                tbl.add_row(t.task_id, req, t.status.value)
            console.print(tbl)
        else:
            for t in self.history[-10:]:
                print(f"  [{t.task_id}] {t.original_request[:50]} — {t.status.value}")

    # ── task processing ──────────────────────────────────────

    def _process(self, user_input: str) -> None:
        if console:
            console.print("[blue]📋 Planning…[/]")

        raw = self.llm.plan_task(user_input)
        plan = ResponseParser.parse_plan(raw)

        if plan is None:
            # Not a task — just a conversational reply
            if console:
                console.print(Panel(raw, title="💬 ATLAS", border_style="blue"))
            else:
                print(f"\n💬 {raw}")
            return

        plan.original_request = user_input

        # Show plan
        if console:
            tree = Tree(f"[cyan]📋 {plan.original_request}[/]")
            for s in plan.steps:
                tree.add(
                    f"[yellow]Step {s.step_number}:[/] {s.description} "
                    f"[dim]({s.tool_name})[/]"
                )
            console.print(tree)
            console.print("[bold]Execute? [Y/n]:[/] ", end="")
        else:
            for s in plan.steps:
                print(
                    f"  Step {s.step_number}: {s.description} ({s.tool_name})"
                )
            print("Execute? [Y/n]: ", end="")

        try:
            ans = input().strip().lower()
        except (EOFError, KeyboardInterrupt):
            print("\nCancelled.")
            return

        if ans in ("n", "no", "nie"):
            print("Cancelled.")
            return

        result = self.engine.run(plan)
        self.history.append(result)

    # ── main loop ────────────────────────────────────────────

    def run(self) -> None:
        self._banner()

        # Connection test
        if console:
            console.print("[dim]Connecting to Ollama…[/]")
        probe = self.llm.chat("Reply with one word: OK")
        if "OLLAMA_ERROR" in probe:
            print(f"⚠️ {probe}")
            print("Make sure Ollama is running:  ollama serve")
            return
        if console:
            console.print("[green]✅ Ollama connected[/]")
        self.llm.reset()

        while self._running:
            try:
                if console:
                    console.print("\n[bold green]🤖 ATLAS>[/] ", end="")
                else:
                    print("\n🤖 ATLAS> ", end="")

                try:
                    user_input = input().strip()
                except (EOFError, KeyboardInterrupt):
                    print("\n👋 Goodbye!")
                    break

                if not user_input:
                    continue

                cmd = user_input.lower()
                if cmd in ("/exit", "/quit", "/q", "exit", "quit"):
                    print("👋 Goodbye!")
                    break
                if cmd == "/help":
                    self._help()
                    continue
                if cmd == "/tools":
                    self._show_tools()
                    continue
                if cmd == "/status":
                    self._show_status()
                    continue
                if cmd == "/history":
                    self._show_history()
                    continue
                if cmd == "/clear":
                    self.llm.reset()
                    print("Conversation cleared.")
                    continue

                self._process(user_input)

            except KeyboardInterrupt:
                print("\nInterrupted — type /exit to quit.")
            except Exception as exc:
                logger.error("Unexpected error: %s", exc, exc_info=True)
                print(f"❌ Error: {exc}")

        self.tools.cleanup()
        logger.info("ATLAS shut down.")


# ════════════════════════════════════════════════════════════════
#  ENTRY POINT
# ════════════════════════════════════════════════════════════════


def main() -> None:
    import argparse

    ap = argparse.ArgumentParser(
        description="ATLAS — AI Task & Automation System",
    )
    ap.add_argument(
        "--model",
        default="jobautomation/OpenEuroLLM-Polish:latest",
        help="Ollama model name",
    )
    ap.add_argument(
        "--task",
        default=None,
        help="Run a single task and exit",
    )
    args = ap.parse_args()

    agent = AtlasAgent(model=args.model)

    if args.task:
        agent._process(args.task)
        agent.tools.cleanup()
    else:
        agent.run()


if __name__ == "__main__":
    main()